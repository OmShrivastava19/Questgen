# backend/app/services/export_service.py

import io
import logging
from docx import Document
from docx.shared import Inches
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet

logger = logging.getLogger(__name__)

class ExportService:
    """
    A service to handle the exporting of question papers to different formats
    like PDF and DOCX.
    """

    def export_to_pdf(self, title: str, questions: list, answer_key: list) -> io.BytesIO:
        """
        Generates a PDF document from a list of questions and an answer key.

        Returns:
            io.BytesIO: A byte stream containing the PDF data.
        """
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []

        # Title
        story.append(Paragraph(title, styles['h1']))
        story.append(Spacer(1, 12))

        # Questions
        for i, q in enumerate(questions):
            q_text = f"{i + 1}. {q.get('question', '')}"
            story.append(Paragraph(q_text, styles['BodyText']))
            
            if q.get('type') == 'mcq' and 'options' in q:
                for j, option in enumerate(q['options']):
                    opt_text = f"   {'ABCD'[j]}) {option}"
                    story.append(Paragraph(opt_text, styles['BodyText']))
            story.append(Spacer(1, 12))

        # Answer Key
        story.append(Paragraph("Answer Key", styles['h2']))
        story.append(Spacer(1, 12))
        for i, ans in enumerate(answer_key):
            ans_text = f"{i + 1}. {ans.get('answer', 'N/A')}"
            story.append(Paragraph(ans_text, styles['BodyText']))

        doc.build(story)
        buffer.seek(0)
        logger.info("Successfully created PDF document in memory.")
        return buffer

    def export_to_docx(self, title: str, questions: list, answer_key: list) -> io.BytesIO:
        """
        Generates a DOCX document from a list of questions and an answer key.

        Returns:
            io.BytesIO: A byte stream containing the DOCX data.
        """
        document = Document()
        document.add_heading(title, level=1)

        # Questions
        for i, q in enumerate(questions):
            document.add_paragraph(f"{i + 1}. {q.get('question', '')}", style='List Number')
            
            if q.get('type') == 'mcq' and 'options' in q:
                for j, option in enumerate(q['options']):
                    p = document.add_paragraph(f"{'ABCD'[j]}) {option}")
                    p.paragraph_format.left_indent = Inches(0.5)

        # Answer Key
        document.add_heading("Answer Key", level=2)
        for i, ans in enumerate(answer_key):
            document.add_paragraph(f"{i + 1}. {ans.get('answer', 'N/A')}", style='List Number')

        buffer = io.BytesIO()
        document.save(buffer)
        buffer.seek(0)
        logger.info("Successfully created DOCX document in memory.")
        return buffer
