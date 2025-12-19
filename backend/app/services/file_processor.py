# services/file_processor.py

import io
import logging
import os
import re
import unicodedata
from typing import List, Dict, Union, BinaryIO, Tuple

import docx
import pdfplumber
import nltk
from nltk.corpus import stopwords
from nltk.tokenize import sent_tokenize, word_tokenize

# --- Setup Logging ---
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# --- Download necessary NLTK data ---
try:
    nltk.data.find('tokenizers/punkt')
    nltk.data.find('corpora/stopwords')
except LookupError:
    logger.warning("NLTK data not found. Downloading 'punkt' and 'stopwords'...")
    nltk.download('punkt', quiet=True)
    nltk.download('stopwords', quiet=True)

# --- Main TextProcessor Class ---

class TextProcessor:
    """
    A comprehensive service to process uploaded text files (PDF, DOCX)
    for consumption by an AI question generation model.

    This class handles file validation, text extraction, cleaning,
    chunking, and key concept identification.
    """

    def __init__(self, max_file_size_mb: int = 10, chunk_size: int = 512):
        """
        Initializes the TextProcessor.

        Args:
            max_file_size_mb (int): The maximum allowed file size in megabytes.
            chunk_size (int): The target number of tokens for each text chunk.
        """
        self.max_file_size_bytes = max_file_size_mb * 1024 * 1024
        self.allowed_extensions = {".pdf", ".docx"}
        self.chunk_size = chunk_size
        self.stop_words = set(stopwords.words('english'))
        logger.info(
            f"TextProcessor initialized. Max file size: {max_file_size_mb}MB, "
            f"Chunk size: {self.chunk_size} tokens."
        )

    def _validate_file(self, file: BinaryIO, filename: str) -> None:
        """
        Validates a single file based on its extension and size.

        Args:
            file (BinaryIO): The file stream object.
            filename (str): The name of the file.

        Raises:
            ValueError: If the file type or size is invalid.
        """
        # Validate file extension
        ext = os.path.splitext(filename)[1].lower()
        if ext not in self.allowed_extensions:
            raise ValueError(
                f"Invalid file type for '{filename}'. "
                f"Allowed types are: {', '.join(self.allowed_extensions)}"
            )
        
        # Validate file size
        file.seek(0, os.SEEK_END)
        file_size = file.tell()
        file.seek(0)
        if file_size > self.max_file_size_bytes:
            raise ValueError(
                f"File '{filename}' exceeds the maximum size of "
                f"{self.max_file_size_bytes / (1024*1024)}MB."
            )
        logger.info(f"File '{filename}' passed validation.")

    def _extract_text_from_pdf(self, file_stream: BinaryIO) -> str:
        """
        Extracts text from a PDF file stream using pdfplumber.
        Handles complex layouts and tables gracefully.
        """
        text = ""
        try:
            with pdfplumber.open(file_stream) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            logger.info("Successfully extracted text from PDF.")
            return text
        except Exception as e:
            logger.error(f"Failed to process PDF file: {e}", exc_info=True)
            raise IOError("Could not extract text from the PDF. It may be corrupted or image-based.")

    def _extract_text_from_docx(self, file_stream: BinaryIO) -> str:
        """
        Extracts text from a DOCX file stream using python-docx.
        Note: This does not support the older .doc format.
        """
        try:
            document = docx.Document(file_stream)
            text = "\n".join([para.text for para in document.paragraphs])
            logger.info("Successfully extracted text from DOCX.")
            return text
        except Exception as e:
            logger.error(f"Failed to process DOCX file: {e}", exc_info=True)
            raise IOError("Could not extract text from the DOCX file. It may be corrupted.")

    def _clean_text(self, text: str) -> str:
        """
        Performs text cleaning and preprocessing.
        
        - Normalizes unicode characters.
        - Removes extra whitespace, newlines, and tabs.
        - Fixes common text artifacts.
        """
        # Normalize unicode characters to handle accents and special chars
        text = unicodedata.normalize('NFKD', text).encode('ascii', 'ignore').decode('utf-8', 'ignore')
        
        # Remove URLs
        text = re.sub(r'http\S+', '', text)
        
        # Replace multiple newlines/tabs with a single space
        text = re.sub(r'[\n\t]+', ' ', text)
        
        # Remove extra whitespace
        text = re.sub(r'\s{2,}', ' ', text).strip()
        
        logger.info(f"Text cleaned. Final length: {len(text)} characters.")
        return text

    def split_into_chunks(self, text: str) -> List[str]:
        """
        Splits the preprocessed text into meaningful chunks of a specified token size.
        Uses NLTK sentence tokenization for better sentence boundary detection.
        """
        if not text:
            return []

        # Use NLTK sentence tokenization
        sentences = sent_tokenize(text)
        
        chunks = []
        current_chunk_tokens = []

        for sentence in sentences:
            # Tokenize the sentence into words
            sentence_tokens = word_tokenize(sentence)

            if len(current_chunk_tokens) + len(sentence_tokens) <= self.chunk_size:
                current_chunk_tokens.extend(sentence_tokens)
            else:
                # Add the current chunk if it's not empty
                if current_chunk_tokens:
                    chunks.append(" ".join(current_chunk_tokens))
                
                # Start a new chunk with the current sentence
                # If a sentence is longer than the chunk size, it will be its own chunk.
                current_chunk_tokens = sentence_tokens

        # Add the last remaining chunk
        if current_chunk_tokens:
            chunks.append(" ".join(current_chunk_tokens))
        
        logger.info(f"Text split into {len(chunks)} chunks of approximately {self.chunk_size} tokens each.")
        return chunks

    def extract_key_concepts(self, text: str, num_concepts: int = 10) -> List[str]:
        """
        Extracts key concepts and topics from the text using basic NLP techniques.
        Identifies significant words and phrases without relying on SpaCy.
        """
        if not text:
            return []
            
        # Tokenize the text
        words = word_tokenize(text.lower())
        
        # Remove stop words and single-character words
        filtered_words = [
            word for word in words 
            if word.lower() not in self.stop_words and len(word) > 2 and word.isalnum()
        ]
        
        # Get the most frequent words
        word_counts = {}
        for word in filtered_words:
            word_counts[word] = word_counts.get(word, 0) + 1
            
        # Sort by frequency, descending
        sorted_words = sorted(word_counts.items(), key=lambda item: item[1], reverse=True)
        
        # Return the top N unique words
        top_concepts = [word for word, count in sorted_words[:num_concepts]]
        
        logger.info(f"Extracted top {len(top_concepts)} key concepts.")
        return top_concepts

    def process_files(self, files: List[Tuple[BinaryIO, str]]) -> Dict[str, Dict]:
        """
        Main method to process a list of uploaded files simultaneously.

        Args:
            files (List[Tuple[BinaryIO, str]]): A list of tuples, where each
                tuple contains a file stream and its filename.

        Returns:
            Dict[str, Dict]: A dictionary where keys are filenames and values
                are dictionaries containing the processed data ('raw_text',
                'cleaned_text', 'chunks', 'key_concepts').
        """
        results = {}
        for file_stream, filename in files:
            try:
                logger.info(f"--- Starting processing for {filename} ---")
                
                # 1. Validate file
                self._validate_file(file_stream, filename)
                
                # 2. Extract text based on file type
                ext = os.path.splitext(filename)[1].lower()
                raw_text = ""
                if ext == ".pdf":
                    raw_text = self._extract_text_from_pdf(file_stream)
                elif ext == ".docx":
                    raw_text = self._extract_text_from_docx(file_stream)

                if not raw_text.strip():
                    raise ValueError("No text could be extracted from the file.")

                # 3. Clean the extracted text
                cleaned_text = self._clean_text(raw_text)
                
                # 4. Split text into chunks for the model
                chunks = self.split_into_chunks(cleaned_text)
                
                # 5. Extract key concepts
                key_concepts = self.extract_key_concepts(cleaned_text)
                
                results[filename] = {
                    "status": "success",
                    "raw_text": raw_text,
                    "cleaned_text": cleaned_text,
                    "chunks": chunks,
                    "key_concepts": key_concepts,
                }
                logger.info(f"--- Finished processing for {filename} ---")

            except (ValueError, IOError, Exception) as e:
                logger.error(f"Failed to process '{filename}': {e}", exc_info=True)
                results[filename] = {
                    "status": "error",
                    "message": str(e)
                }
        return results

# --- Unit Test Examples ---
# This section demonstrates how to test the TextProcessor class.
# In a real project, this would be in a separate `tests/` directory.

import unittest
from unittest.mock import MagicMock

class TestTextProcessor(unittest.TestCase):

    def setUp(self):
        """Set up a TextProcessor instance for all tests."""
        self.processor = TextProcessor(max_file_size_mb=1)

    def test_validate_file_success(self):
        """Test that valid files pass validation."""
        mock_file = io.BytesIO(b"small file content")
        self.assertIsNone(self.processor._validate_file(mock_file, "test.pdf"))
        mock_file.seek(0)
        self.assertIsNone(self.processor._validate_file(mock_file, "test.docx"))

    def test_validate_file_invalid_type(self):
        """Test that invalid file types raise ValueError."""
        mock_file = io.BytesIO(b"content")
        with self.assertRaisesRegex(ValueError, "Invalid file type"):
            self.processor._validate_file(mock_file, "test.txt")

    def test_validate_file_too_large(self):
        """Test that oversized files raise ValueError."""
        large_content = b"a" * (2 * 1024 * 1024) # 2MB
        mock_file = io.BytesIO(large_content)
        with self.assertRaisesRegex(ValueError, "exceeds the maximum size"):
            self.processor._validate_file(mock_file, "large.pdf")

    def test_clean_text(self):
        """Test the text cleaning functionality."""
        raw_text = "  This is a   test.\n\nIt has extra    spaces and\ttabs. Visit http://example.com.  "
        expected_text = "This is a test. It has extra spaces and tabs. Visit ."
        self.assertEqual(self.processor._clean_text(raw_text), expected_text)

    def test_split_into_chunks(self):
        """Test text chunking logic."""
        # This text is designed to have more than 512 tokens if chunk_size is small
        processor = TextProcessor(chunk_size=10)
        text = "This is the first sentence. This is the second sentence, it is a bit longer. Finally, a third one."
        chunks = processor.split_into_chunks(text)
        self.assertEqual(len(chunks), 2)
        self.assertTrue("first sentence" in chunks[0])
        self.assertTrue("second sentence" in chunks[0])
        self.assertTrue("third one" in chunks[1])
        self.assertLessEqual(len(chunks[0].split()), 10)

    def test_extract_key_concepts(self):
        """Test key concept extraction."""
        text = "Apple Inc. is a technology company headquartered in Cupertino, California. Steve Jobs was its co-founder."
        concepts = self.processor.extract_key_concepts(text)
        self.assertIn("Apple Inc.", concepts)
        self.assertIn("Steve Jobs", concepts)
        self.assertIn("Cupertino", concepts)
        self.assertIn("California", concepts)
        self.assertIn("technology company", concepts)

    def test_process_files_unsupported_doc(self):
        """Test handling of unsupported .doc files."""
        # python-docx does not support .doc, so we expect an error.
        # This test simulates that by passing a file with a .docx extension
        # but in a format that python-docx can't read (e.g., an empty file).
        mock_file = io.BytesIO(b"")
        files_to_process = [(mock_file, "document.docx")]
        results = self.processor.process_files(files_to_process)
        self.assertEqual(results["document.docx"]["status"], "error")
        self.assertIn("Could not extract text", results["document.docx"]["message"])

if __name__ == '__main__':
    # Example usage of the TextProcessor
    print("--- Running TextProcessor Example ---")
    
    # Create a dummy DOCX file in memory for demonstration
    from docx import Document
    doc = Document()
    doc.add_paragraph("This is the first paragraph of our test document.")
    doc.add_paragraph("SpaCy is a powerful library for Natural Language Processing. It helps with tokenization and named entity recognition.")
    doc_stream = io.BytesIO()
    doc.save(doc_stream)
    doc_stream.seek(0)

    # Create a dummy text file to simulate another upload
    text_content = b"This is a plain text file that will be rejected."
    text_stream = io.BytesIO(text_content)
    
    # Initialize processor and process files
    processor = TextProcessor()
    files = [
        (doc_stream, "example.docx"),
        (text_stream, "unsupported.txt")
    ]
    
    processed_data = processor.process_files(files)

    # Print results
    import json
    print(json.dumps(processed_data, indent=2))

    # --- Running Unit Tests ---
    print("\n--- Running Unit Tests ---")
    suite = unittest.TestSuite()
    suite.addTest(unittest.makeSuite(TestTextProcessor))
    runner = unittest.TextTestRunner()
    runner.run(suite)
