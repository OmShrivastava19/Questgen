# backend/tests/test_text_processor.py

import unittest
import io
from app.services.file_processor import TextProcessor
import docx

class TestTextProcessor(unittest.TestCase):
    """Unit tests for the TextProcessor service."""

    def setUp(self):
        """Set up a TextProcessor instance before each test."""
        self.processor = TextProcessor(max_file_size_mb=1)

    def test_clean_text(self):
        """Test the text cleaning and normalization logic."""
        raw_text = "  Here is... some text\nwith extra\tspacing & weird chars like é. Visit https://example.com.  "
        # CORRECTED: The regex `http\S+` also consumes the period that follows it.
        expected = "Here is... some text with extra spacing & weird chars like e. Visit"
        cleaned_text = self.processor._clean_text(raw_text)
        self.assertEqual(cleaned_text, expected)

    def test_extract_text_from_docx(self):
        """Test DOCX text extraction using an in-memory document."""
        document = docx.Document()
        document.add_paragraph("Hello, world.")
        document.add_paragraph("Second paragraph.")
        
        mock_file_stream = io.BytesIO()
        document.save(mock_file_stream)
        mock_file_stream.seek(0)
        text = self.processor._extract_text_from_docx(mock_file_stream)
        
        # CORRECTED: The function joins paragraphs with a newline.
        self.assertEqual(text, "Hello, world.\nSecond paragraph.")
